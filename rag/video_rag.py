import os
import re
import PyPDF2
from docx import Document
import pytesseract
from pdf2image import convert_from_path
from langchain import OpenAI
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain import LLMChain
from langchain.prompts.prompt import PromptTemplate
from typing import List 
from fastapi import FastAPI, File, UploadFile, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse 
import mysql.connector as msql
import chromadb 

tesseract_cmd_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
pytesseract.pytesseract.tesseract_cmd = tesseract_cmd_path

app = FastAPI()

origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

mydb = msql.connector.connect(
    host = "my_sql_host",
    user = "tyson",
    password ="t6NDFLNKIWZE",
    database = "tysondb"
)

mycursor = mydb.cursor()

vectordb = Chroma(persist_directory = "embeddings")

base_dir = os.getcwd()
if not os.path.exists(os.path.join(base_dir, 'Documents')):
    os.mkdir(os.path.join(base_dir, 'Documents'))
pdf_dir = os.path.join(base_dir, 'Documents')

# Cleaning extracted text 
def clean_text(pdf_path):
    cleaned_text = re.sub(r'[^a-zA-Z0-9\s]','',text)
    cleaned_text = ''.join(cleaned_text.split())
    return cleaned_text

def extract_text_from_pdf(pdf_path):
    texts = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=10)

    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page_num in range(pdf_reader.numPages):
            text = pdf_reader.getPage(page_num).extractText()
            if text.strip() == "":  # Employ Tesseract if PyPDF2 fails
                images = convert_from_path(pdf_path)
                doc = []
                for page_number, image in enumerate(images, start=1):
                    image = image.convert('L')
                    text = pytesseract.image_to_string(image)
                    doc.append(Document(page_content=text, metadata={"source": pdf_path, "page": page_number, "extension": "pdf"}))
                texts.extend(text_splitter.split_documents(doc))
            else:
                doc = Document(page_content=text, metadata={"source": pdf_path, "page": page_num, "extension": "pdf"})
                texts.extend(text_splitter.split_documents([doc]))
    return texts

def extract_text_from_word(docx_path):
    texts = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=10)

    doc = Document(docx_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    doc = Document(page_content=text, metadata={"source": docx_path, "extension": "docx"})
    texts.extend(text_splitter.split_documents([doc]))
    return texts

def process_folder(folder_path):
    extracted_data = []

    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            texts = extract_text_from_pdf(file_path)
            extracted_data.extend(texts)
        elif filename.endswith(".docx"):
            texts = extract_text_from_word(file_path)
            extracted_data.extend(texts)
    return extracted_data


extracted_text_chunks = process_folder(folder_path)

mydb = msql.connector.connect(
    host = "my_sql_host",
    user = "tyson",
    password ="t6NDFLNKIWZE",
    database = "tysondb"
)

mycursor = mydb.cursor()

mycursor.execute("""
                 CREATE TABLE IF NOT EXISTS files (
                 id INT AUTO_INCREMENT PRIMARY KEY,
                 filename VARCHAR(255) NOT NULL UNIQUE
                 )
                 """)

mycursor.execute("""
                 CREATE TABLE IF NOT EXISTS extracted_text (
                 id INT AUTO_INCREMENT PRIMARY KEY,
                 file_id INT NOT NULL,
                 page_content TEXT NOT NULL,
                 page INT,
                 extension VARCHAR(10),
                 FOREIGN KEY (file_id) REFERENCES files(id)
                 )
                 
                 """)

def insert_files_in_tables(filename):
    folder_path = os.path.join(os.getcwd(), 'Documents')
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
    # Insert file information into the 'files' table
    sql = "INSERT INTO files (filename) VALUES (%s)"
    val = (filename,)
    mycursor.execute(sql, val)
    mydb.commit()  # Commit changes to the database

    # Insert extracted text chunks into the 'extracted_text' table
    for chunk in extracted_text_chunks:
        if chunk.metadata['source'] == file_path:  # Match chunks to the current file
            sql = "INSERT INTO extracted_text (file_id, page_content, page, extension) VALUES (%s, %s, %s, %s)"
            val = (mycursor.lastrowid, chunk.page_content, chunk.page, chunk.metadata['extension'])
            mycursor.execute(sql, val)
            mydb.commit()